from flask import Blueprint, render_template, request, redirect, url_for, flash, Response
from flask_login import login_required, current_user
from ..utils.database import get_db_connection, get_user_clients, get_templates_for_client, get_conversion_prompts_for_client
from ..utils.document import process_docx, process_text_input
from ..utils.conversion import convert_content
from ..utils.docx_builder import create_reformatted_docx
from docx import Document
from docx.shared import Pt
import json
from io import BytesIO
import logging
import re

logger = logging.getLogger(__name__)

main_bp = Blueprint('main', __name__)

@main_bp.route('/', methods=['GET', 'POST'])
@login_required
def index():
    clients = get_user_clients(current_user.id)
    selected_client = request.args.get('client_id', '') if request.method == 'GET' else request.form.get('client', '')
    templates = get_templates_for_client(selected_client, current_user.id)
    conversion_prompts = get_conversion_prompts_for_client(selected_client, current_user.id)

    template_prompt = ''
    conversion_prompt = ''
    selected_template = ''
    conversion_prompt_id = ''

    # Filter templates and conversion prompts for the selected client
    filtered_templates = [
        template for template in templates
        if (selected_client and (template['client_id'] is None or template['client_id'] == selected_client)) or
           (not selected_client and template['client_id'] is None)
    ]
    filtered_conversion_prompts = [
        prompt for prompt in conversion_prompts
        if (selected_client and (prompt['client_id'] is None or prompt['client_id'] == selected_client)) or
           (not selected_client and prompt['client_id'] is None)
    ]

    if request.method == 'POST':
        action = request.form.get('action')
        selected_template = request.form.get('template', '')
        template_prompt = request.form.get('template_prompt', '')
        conversion_prompt = request.form.get('conversion_prompt', '')
        conversion_prompt_id = request.form.get('conversion_prompt_id', '')

        if action == 'select_client':
            return redirect(url_for('main.index', client_id=selected_client))

        elif action == 'convert':
            if not selected_template:
                flash('Please select a template', 'danger')
                return redirect(url_for('main.index', client_id=selected_client))
            if not template_prompt:
                flash('Template prompt is required', 'danger')
                return redirect(url_for('main.index', client_id=selected_client))

            try:
                # Get template file and prompt for styling and structuring
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT template_file, template_prompt_id, p.content AS template_prompt_content "
                    "FROM templates t "
                    "LEFT JOIN prompts p ON t.template_prompt_id = p.id "
                    "WHERE t.id = %s AND t.user_id = %s",
                    (selected_template, current_user.id)
                )
                template_data = cur.fetchone()
                template_file = template_data[0]
                template_prompt_id = template_data[1]
                template_prompt_content = template_data[2]
                cur.close()
                conn.close()

                if not template_file:
                    flash('Template file not found. Please ensure the selected template has an associated file.', 'danger')
                    return redirect(url_for('main.index', client_id=selected_client))
                if not template_prompt_content:
                    flash('Template prompt content not found. Please ensure the selected template has an associated prompt.', 'danger')
                    return redirect(url_for('main.index', client_id=selected_client))

                logger.info(f"Using template file for template ID {selected_template} (length: {len(template_file)} bytes)")

                # Parse expected sections from the template prompt
                expected_sections = []
                section_pattern = r"\*\*Section:\s*([^\*]+)\*\*\s*- \*\*Purpose\*\*:\s*This section represents\s*([^\s]+)\s*content"
                for match in re.finditer(section_pattern, template_prompt_content):
                    section_name = match.group(1).strip()
                    section_key = match.group(2).strip()
                    expected_sections.append((section_name.lower(), section_key))

                # Process source content
                source_file = request.files.get('source_file')
                if source_file and source_file.filename.endswith('.docx'):
                    # For .docx files, extract content directly
                    doc = Document(source_file)
                    sections = []
                    current_section = None
                    raw_content = []

                    # Extract all paragraphs and tables
                    for element in doc.element.body:
                        if element.tag.endswith('p'):  # Paragraph
                            para = doc.paragraphs[len(raw_content)]
                            text = para.text.strip()
                            if text:
                                raw_content.append({"type": "paragraph", "text": text, "runs": para.runs})
                        elif element.tag.endswith('tbl'):  # Table
                            table = doc.tables[len([e for e in raw_content if e["type"] == "table"])]
                            table_content = []
                            for row in table.rows:
                                row_content = [cell.text.strip() for cell in row.cells]
                                table_content.append(row_content)
                            raw_content.append({"type": "table", "content": table_content})

                    # Map content to expected sections using template prompt
                    structured_content = {"sections": {}}
                    used_content_indices = set()

                    for section_name, section_key in expected_sections:
                        best_match_idx = -1
                        best_match_score = 0
                        for idx, item in enumerate(raw_content):
                            if idx in used_content_indices:
                                continue
                            if item["type"] == "paragraph":
                                text = item["text"].lower()
                                # Simple scoring based on header similarity
                                if section_name in text:
                                    score = 1.0  # Exact match
                                else:
                                    # Approximate match based on keywords
                                    keywords = section_name.split()
                                    score = sum(1 for keyword in keywords if keyword in text) / len(keywords)
                                if score > best_match_score:
                                    best_match_score = score
                                    best_match_idx = idx

                        if best_match_idx >= 0:
                            # Found a matching paragraph, start collecting content until the next section
                            content = []
                            idx = best_match_idx
                            while idx < len(raw_content):
                                item = raw_content[idx]
                                if idx in used_content_indices:
                                    idx += 1
                                    continue
                                if item["type"] == "paragraph":
                                    text = item["text"]
                                    # Check if this paragraph matches another section header
                                    is_new_section = False
                                    for other_section_name, _ in expected_sections:
                                        if other_section_name != section_name and other_section_name in text.lower():
                                            is_new_section = True
                                            break
                                    if is_new_section:
                                        break
                                    content.append(text)
                                    used_content_indices.add(idx)
                                elif item["type"] == "table":
                                    # Tables are handled separately
                                    break
                                idx += 1
                            structured_content["sections"][section_key] = content
                            used_content_indices.add(best_match_idx)

                    # Handle tables separately
                    tables = []
                    for idx, item in enumerate(raw_content):
                        if idx in used_content_indices:
                            continue
                        if item["type"] == "table":
                            tables.append(item["content"])
                            used_content_indices.add(idx)
                    if tables:
                        structured_content["sections"]["tables"] = tables

                    # Fill missing sections with empty lists
                    for _, section_key in expected_sections:
                        if section_key not in structured_content["sections"]:
                            structured_content["sections"][section_key] = []

                else:
                    # For non-.docx sources, use LLM to interpret content
                    source_text = request.form.get('source_text', '')
                    if not source_text:
                        flash('Please upload a .docx file or provide text input', 'danger')
                        return redirect(url_for('main.index', client_id=selected_client))
                    content = process_text_input(source_text)
                    structured_content = convert_content(content, template_prompt, conversion_prompt)

                # Apply styles with python-docx
                output_file = create_reformatted_docx(structured_content, template_file)

                # Return the file for immediate download
                return Response(
                    output_file,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': 'attachment; filename=reformatted_document.docx'}
                )
            except TypeError as e:
                flash(f"Conversion failed due to invalid input types: {str(e)}. Please ensure the template and conversion prompts are correctly formatted.", 'danger')
                return redirect(url_for('main.index', client_id=selected_client))
            except Exception as e:
                flash(f"Conversion failed: {str(e)}. Please check the template and conversion prompts and try again.", 'danger')
                return redirect(url_for('main.index', client_id=selected_client))

    return render_template(
        'index.html',
        clients=clients,
        selected_client=selected_client,
        templates=filtered_templates,
        conversion_prompts=filtered_conversion_prompts,
        template_prompt=template_prompt,
        conversion_prompt=conversion_prompt,
        selected_template=selected_template,
        conversion_prompt_id=conversion_prompt_id
    )

@main_bp.route('/load_client', methods=['POST'])
def load_client():
    client_id = request.form.get('client_id', '')
    template_id = request.form.get('template_id', '')
    prompt_id = request.form.get('prompt_id', '')

    if template_id:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT p.content AS template_prompt, cp.content AS conversion_prompt "
            "FROM templates t "
            "JOIN prompts p ON t.template_prompt_id = p.id "
            "LEFT JOIN template_prompt_associations tpa ON t.id = tpa.template_id "
            "LEFT JOIN prompts cp ON tpa.conversion_prompt_id = cp.id "
            "WHERE t.id = %s AND t.user_id = %s",
            (template_id, current_user.id)
        )
        result = cur.fetchone()
        cur.close()
        conn.close()
        if result:
            return {'prompt': result[0], 'conversion': result[1] if result[1] else ''}
        return {'prompt': '', 'conversion': ''}

    if prompt_id:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT content FROM prompts WHERE id = %s AND user_id = %s",
            (prompt_id, current_user.id)
        )
        result = cur.fetchone()
        cur.close()
        conn.close()
        if result:
            return {'prompt': result[0]}
        return {'prompt': ''}

    return {'prompt': '', 'conversion': ''}