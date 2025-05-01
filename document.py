from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
import json
import os
from database import load_ai_prompt

AI_API_URL = os.environ.get('AI_API_URL', 'https://api.openai.com/v1/chat/completions')
API_KEY = os.environ.get('API_KEY', 'your-api-key')

def extract_content_from_docx(file_path):
    try:
        doc = Document(file_path)
        content = {"text": [], "tables": [], "references": []}
        in_references = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if text.lower().startswith("references"):
                    in_references = True
                    continue
                if in_references:
                    content["references"].append(text)
                else:
                    content["text"].append(text)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_data.append(row_data)
            if table_data:
                print(f"Extracted table: {table_data}")
                content["tables"].append(table_data)
        return content
    except Exception as e:
        print(f"Error extracting content from docx: {str(e)}")
        raise

def call_ai_api(content, client_id=None, user_id=None, prompt_name=None):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    text = "\n".join(content["text"])
    tables = json.dumps(content["tables"])
    ai_prompt = load_ai_prompt(client_id, user_id, prompt_name)
    messages = [
        {"role": "system", "content": ai_prompt},
        {
            "role": "user",
            "content": (
                f"Input Text:\n{text}\n\nTables:\n{tables}\n\nOutput format:\n"
                "{\"summary\": \"...\", \"background\": \"...\", \"monograph\": \"...\", "
                "\"real_world\": \"\", \"enclosures\": \"...\", "
                "\"tables\": {\"section_name\": [[\"cell1\", \"cell2\"], [\"cell3\", \"cell4\"]]}, "
                "\"references\": [\"ref1\", \"ref2\", ...]}"
            )
        }
    ]
    payload = {
        "model": "gpt-3.5-turbo",
        "messages": messages,
        "max_tokens": 2000,
        "temperature": 0.7
    }
    try:
        response = requests.post(AI_API_URL, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        raw_content = data["choices"][0]["message"]["content"]
        print(f"Raw AI response: {raw_content[:1000]}...")
        try:
            parsed_content = json.loads(raw_content)
            if not isinstance(parsed_content, dict):
                raise ValueError("AI response is not a JSON object")
            if "tables" in parsed_content and isinstance(parsed_content["tables"], dict):
                for section_name, table_data in parsed_content["tables"].items():
                    while isinstance(table_data, list) and len(table_data) == 1 and isinstance(table_data[0], list):
                        table_data = table_data[0]
                    parsed_content["tables"][section_name] = table_data
                    if not all(isinstance(row, list) for row in table_data):
                        print(f"Invalid table data for {section_name}: {table_data}")
                        parsed_content["tables"][section_name] = []
            print(f"Parsed AI response: {parsed_content}")
            return parsed_content
        except json.JSONDecodeError as e:
            print(f"JSON validation error: {str(e)}")
            return {
                "error": f"Invalid JSON from AI: {str(e)}",
                "summary": "Unable to categorize due to AI response error",
                "background": content["text"][:500] if content["text"] else "",
                "monograph": "",
                "real_world": "",
                "enclosures": "",
                "tables": {},
                "references": content["references"]
            }
    except requests.exceptions.HTTPError as e:
        error_response = response.json() if response else {"error": str(e)}
        print(f"API Error: {response.status_code} - {error_response}")
        return {"error": f"HTTP Error: {str(e)} - {error_response}"}
    except Exception as e:
        print(f"Unexpected Error: {str(e)}")
        return {"error": str(e)}

def add_styled_heading(doc, text, level=1):
    try:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.underline = True if level == 1 else False
        run.font.name = "Arial"
        run.font.size = Pt(14)
        return para
    except Exception as e:
        print(f"Error adding styled heading: {str(e)}")
        raise

def add_styled_text(doc, text, bullet=False):
    try:
        para = doc.add_paragraph(style="List Bullet" if bullet else None)
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        return para
    except Exception as e:
        print(f"Error adding styled text: {str(e)}")
        raise

def add_styled_table(doc, table_data, section_name):
    try:
        if not table_data or not table_data[0] or not any(cell for row in table_data for cell in row):
            print(f"Skipping invalid or empty table for section: {section_name}")
            return None
        max_cols = max(len(row) for row in table_data)
        table_data = [row + [""] * (max_cols - len(row)) for row in table_data]
        print(f"Adding table for {section_name}: {table_data}")
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            print(f"Processing row {i}: {row_data}")
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text or ""
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                print(f"Setting borders for cell at row {i}, col {j}")
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if not tcBorders:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)
        return table
    except Exception as e:
        print(f"Error adding styled table for section {section_name}: {str(e)}")
        raise

def create_reformatted_docx(sections, output_path, drug_name="KRESLADI", client_id=None, user_id=None):
    try:
        default_sections = {
            "summary": "No summary provided",
            "background": "No background provided",
            "monograph": "No monograph provided",
            "real_world": "",
            "enclosures": "No enclosures provided",
            "tables": {},
            "references": []
        }
        sections = {**default_sections, **sections}
        template_path = os.path.join(os.environ.get('UPLOAD_FOLDER', '/tmp'), 'temp_template.docx')
        if load_template(template_path, client_id, user_id):
            print(f"Using template for client {client_id}, user {user_id}: {template_path}")
            doc = Document(template_path)
            def replace_placeholder(paragraph, placeholder, content, preserve_style=True):
                if placeholder.lower() in paragraph.text.lower():
                    if preserve_style:
                        paragraph.text = ""
                        run = paragraph.add_run(content)
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            run.bold = first_run.bold
                            run.underline = first_run.underline
                            run.font.name = first_run.font.name
                            run.font.size = first_run.font.size
                    else:
                        paragraph.text = content
            def add_table_after_placeholder(doc, placeholder, table_data, section_name):
                for i, para in enumerate(doc.paragraphs):
                    if placeholder.lower() in para.text.lower():
                        print(f"Adding table for {section_name} after placeholder: {placeholder}")
                        max_cols = max(len(row) for row in table_data)
                        table_data = [row + [""] * (max_cols - len(row)) for row in table_data]
                        table = doc.add_table(rows=len(table_data), cols=max_cols)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.autofit = True
                        for row_idx, row_data in enumerate(table_data):
                            row = table.rows[row_idx]
                            for col_idx, cell_text in enumerate(row_data):
                                cell = row.cells[col_idx]
                                cell.text = cell_text or ""
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.name = para.runs[0].font.name if para.runs else "Calibri"
                                        r.font.size = para.runs[0].font.size if para.runs else Pt(10)
                        if doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                                    if not tcBorders:
                                        tcBorders = OxmlElement('w:tcBorders')
                                        tcPr.append(tcBorders)
                                        for border_name in ['top', 'left', 'bottom', 'right']:
                                            border = OxmlElement(f'w:{border_name}')
                                            border.set(qn('w:val'), 'single')
                                            border.set(qn('w:sz'), '4')
                                            border.set(qn('w:space'), '0')
                                            border.set(qn('w:color'), 'auto')
                                            tcBorders.append(border)
                        return True
                return False
            for para in doc.paragraphs:
                if "drug name" in para.text.lower():
                    replace_placeholder(para, "drug name", f"{drug_name} (marnetegragene autotemcel)")
                elif "summary" in para.text.lower():
                    replace_placeholder(para, "summary", sections["summary"])
                elif "background" in para.text.lower():
                    replace_placeholder(para, "background", sections["background"])
                elif "monograph" in para.text.lower():
                    replace_placeholder(para, "monograph", sections["monograph"])
                elif "real-world experiences" in para.text.lower() and sections["real_world"].strip():
                    replace_placeholder(para, "real-world experiences", sections["real_world"])
                elif "enclosures" in para.text.lower():
                    replace_placeholder(para, "enclosures", sections["enclosures"])
                elif "references" in para.text.lower():
                    replace_placeholder(para, "references", "\n".join([f"{i}. {ref}" for i, ref in enumerate(sections["references"], 1)]))
            for section_name, table_data in sections["tables"].items():
                if not add_table_after_placeholder(doc, section_name, table_data, section_name):
                    print(f"No placeholder found for table: {section_name}, appending at end")
                    para = doc.add_paragraph(section_name)
                    add_styled_table(doc, table_data, section_name)
        else:
            print(f"No template found for client {client_id}, user {user_id}, using default formatting")
            doc = Document()
            add_styled_heading(doc, f"{drug_name} (marnetegragene autotemcel)", level=1)
            add_styled_heading(doc, "Summary", level=1)
            for line in sections["summary"].split("\n"):
                if line.strip():
                    add_styled_text(doc, line, bullet=True)
            add_styled_heading(doc, "Background Information on Leukocyte Adhesion Deficiency (LAD-I)", level=1)
            for line in sections["background"].split("\n"):
                if line.strip():
                    add_styled_text(doc, line)
            add_styled_heading(doc, "Product Monograph", level=1)
            for line in sections["monograph"].split("\n"):
                if line.strip():
                    add_styled_text(doc, line, bullet=True)
            if sections["real_world"].strip():
                add_styled_heading(doc, "Real-World Experiences with KRESLADI", level=1)
                for line in sections["real_world"].split("\n"):
                    if line.strip():
                        add_styled_text(doc, line)
            print(f"Processing tables: {sections['tables']}")
            for section_name, table_data in sections["tables"].items():
                add_styled_heading(doc, section_name, level=2)
                add_styled_table(doc, table_data, section_name)
            add_styled_heading(doc, "Figures", level=1)
            add_styled_text(doc, "Insert Figure 1: Study Administration and Treatment here")
            add_styled_text(doc, "Insert Figure 2: Incidence of Infection-related Hospitalizations here")
            add_styled_heading(doc, "Enclosures", level=1)
            for line in sections["enclosures"].split("\n"):
                if line.strip():
                    add_styled_text(doc, line, bullet=True)
            add_styled_heading(doc, "References", level=1)
            for i, ref in enumerate(sections["references"], 1):
                add_styled_text(doc, f"{i}. {ref}", bullet=False)
        doc.save(output_path)
    except Exception as e:
        print(f"Error creating reformatted docx: {str(e)}")
        raise