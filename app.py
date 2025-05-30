from flask import Flask, request, send_file, render_template, jsonify, redirect, url_for, flash, session
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import json
import os
import psycopg2
from psycopg2.extras import Json
from werkzeug.utils import secure_filename
import bcrypt
from authlib.integrations.flask_client import OAuth
import secrets
import logging
import concurrent.futures
from tempfile import TemporaryDirectory
from hashlib import md5
from io import BytesIO

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key')
app.config['GOOGLE_CLIENT_ID'] = os.environ.get('GOOGLE_CLIENT_ID')
app.config['GOOGLE_CLIENT_SECRET'] = os.environ.get('GOOGLE_CLIENT_SECRET')
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL')

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Flask-Login setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# OAuth setup
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=app.config['GOOGLE_CLIENT_ID'],
    client_secret=app.config['GOOGLE_CLIENT_SECRET'],
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'}
)

# Database setup
DATABASE_URL = os.environ.get('DATABASE_URL')
def get_db_connection():
    conn = psycopg2.connect(DATABASE_URL, sslmode='require')
    return conn

def init_db():
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                email VARCHAR(255) UNIQUE NOT NULL,
                password_hash VARCHAR(255),
                google_id VARCHAR(255) UNIQUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            CREATE TABLE IF NOT EXISTS settings (
                id SERIAL PRIMARY KEY,
                user_id INTEGER REFERENCES users(id),
                client_id VARCHAR(50) NOT NULL,
                prompt JSONB,
                prompt_name VARCHAR(255),
                template BYTEA,
                template_name VARCHAR(255),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                CONSTRAINT settings_unique_user_client_prompt_template UNIQUE (user_id, client_id, prompt_name, template_name)
            );
            CREATE INDEX IF NOT EXISTS idx_client_id ON settings(client_id);
        """)
        conn.commit()
        cur.close()
        conn.close()
        logger.info("Initialized database schema")
    except Exception as e:
        logger.error(f"Error initializing database: {str(e)}")
        raise

with app.app_context():
    init_db()

# User model
class User(UserMixin):
    def __init__(self, id, email, google_id=None):
        self.id = id
        self.email = email
        self.google_id = google_id

@login_manager.user_loader
def load_user(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id, email, google_id FROM users WHERE id = %s", (user_id,))
        user = cur.fetchone()
        cur.close()
        conn.close()
        if user:
            return User(user[0], user[1], user[2])
        return None
    except Exception as e:
        logger.error(f"Error loading user: {str(e)}")
        return None

# Authentication routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("SELECT id, email, password_hash FROM users WHERE email = %s", (email,))
            user = cur.fetchone()
            cur.close()
            conn.close()
            if user and bcrypt.checkpw(password.encode('utf-8'), user[2].encode('utf-8')):
                login_user(User(user[0], user[1]))
                return redirect(url_for('index'))
            flash('Invalid email or password', 'danger')
        except Exception as e:
            logger.error(f"Error during login: {str(e)}")
            flash('Login failed', 'danger')
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("SELECT id FROM users WHERE email = %s", (email,))
            if cur.fetchone():
                flash('Email already registered', 'danger')
                cur.close()
                conn.close()
                return render_template('register.html')
            password_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cur.execute("INSERT INTO users (email, password_hash) VALUES (%s, %s) RETURNING id", (email, password_hash))
            user_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            conn.close()
            login_user(User(user_id, email))
            return redirect(url_for('index'))
        except Exception as e:
            logger.error(f"Error during registration: {str(e)}")
            flash('Registration failed', 'danger')
    return render_template('register.html')

@app.route('/google_login')
def google_login():
    nonce = secrets.token_urlsafe(16)
    session['nonce'] = nonce
    redirect_uri = url_for('google_auth', _external=True)
    return google.authorize_redirect(redirect_uri, nonce=nonce)

@app.route('/google_auth')
def google_auth():
    try:
        token = google.authorize_access_token()
        if not token:
            raise ValueError("No token received from Google")
        nonce = session.pop('nonce', None)
        if not nonce:
            raise ValueError("Nonce not found in session")
        user_info = google.parse_id_token(token, nonce=nonce)
        if not user_info:
            raise ValueError("Failed to parse user info from token")
        google_id = user_info.get('sub')
        email = user_info.get('email')
        if not google_id or not email:
            raise ValueError(f"Missing user info: google_id={google_id}, email={email}")
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id, email FROM users WHERE google_id = %s OR email = %s", (google_id, email))
        user = cur.fetchone()
        if user:
            login_user(User(user[0], user[1], google_id))
        else:
            cur.execute("INSERT INTO users (email, google_id) VALUES (%s, %s) RETURNING id", (email, google_id))
            user_id = cur.fetchone()[0]
            conn.commit()
            login_user(User(user_id, email, google_id))
        cur.close()
        conn.close()
        return redirect(url_for('index'))
    except Exception as e:
        logger.error(f"Error during Google auth: {str(e)}")
        flash(f'Google login failed: {str(e)}', 'danger')
        return redirect(url_for('login'))

@app.route('/logout')
@login_required
def logout():
    session.pop('selected_client', None)
    session.pop('selected_prompt', None)
    session.pop('selected_template', None)
    logout_user()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))

# Client creation route
@app.route('/create_client', methods=['GET', 'POST'])
@login_required
def create_client():
    clients = get_user_clients(current_user.id)
    selected_client = request.args.get('selected_client', '')
    prompts = []
    templates = []

    if selected_client:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND client_id = %s AND prompt IS NOT NULL AND template_name IS NULL",
            (current_user.id, selected_client)
        )
        prompts = [{'prompt_name': row[0], 'prompt_content': row[1] or ''} for row in cur.fetchall()]
        cur.execute(
            "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND client_id = %s AND template IS NOT NULL",
            (current_user.id, selected_client)
        )
        templates = [{'template_name': row[0], 'prompt_name': row[1]} for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Loaded prompts for client {selected_client}: {prompts}")
        logger.info(f"Loaded templates for client {selected_client}: {templates}")

    if request.method == 'POST':
        action = request.form.get('action')
        client_id = request.form.get('client_id', '').strip()

        if action == 'create':
            if not client_id:
                flash('Client ID cannot be empty', 'danger')
                return redirect(url_for('create_client'))
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT id FROM settings WHERE user_id = %s AND client_id = %s",
                    (current_user.id, client_id)
                )
                if cur.fetchone():
                    flash(f'Client "{client_id}" already exists', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('create_client'))
                cur.execute(
                    "INSERT INTO settings (user_id, client_id) VALUES (%s, %s)",
                    (current_user.id, client_id)
                )
                conn.commit()
                cur.close()
                conn.close()
                flash(f'Client {client_id} created successfully', 'success')
                return redirect(url_for('create_client', selected_client=client_id))
            except Exception as e:
                logger.error(f"Error creating client: {str(e)}")
                flash(f'Failed to create client: {str(e)}', 'danger')
                return redirect(url_for('create_client'))

    return render_template('create_client.html', clients=clients, selected_client=selected_client, prompts=prompts, templates=templates)

# Prompt creation/editing route
@app.route('/create_prompt', methods=['GET', 'POST'])
@login_required
def create_prompt():
    clients = get_user_clients(current_user.id)
    selected_client = request.args.get('client_id', '')
    edit_prompt = request.args.get('edit_prompt', '')
    prompts = []
    
    if request.method == 'POST':
        action = request.form.get('action')
        client_id = request.form.get('client_id', '').strip()
        prompt_name = request.form.get('prompt_name', '').strip()
        prompt_content = request.form.get('prompt_content', '').strip()
        original_prompt_name = request.form.get('original_prompt_name', prompt_name).strip()
        
        if action == 'create':
            if not prompt_name:
                flash('Prompt name cannot be empty', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
            if not prompt_content:
                flash('Prompt content cannot be empty', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND prompt_name = %s AND template_name IS NULL",
                    (current_user.id, client_id, prompt_name)
                )
                if cur.fetchone():
                    flash(f'Prompt "{prompt_name}" already exists for this client', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('create_prompt', client_id=client_id))
                save_prompt(prompt_content, client_id, current_user.id, prompt_name)
                flash(f'Prompt "{prompt_name}" created successfully', 'success')
                cur.close()
                conn.close()
                return redirect(url_for('create_prompt', client_id=client_id))
            except Exception as e:
                logger.error(f"Error creating prompt: {str(e)}")
                flash(f'Failed to create prompt: {str(e)}', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
        
        elif action == 'update':
            if not prompt_name:
                flash('Prompt name cannot be empty', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
            if not prompt_content:
                flash('Prompt content cannot be empty', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                if prompt_name != original_prompt_name:
                    cur.execute(
                        "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND prompt_name = %s AND template_name IS NULL",
                        (current_user.id, client_id, prompt_name)
                    )
                    if cur.fetchone():
                        flash(f'Prompt "{prompt_name}" already exists for this client', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('create_prompt', client_id=client_id))
                cur.execute(
                    "UPDATE settings SET prompt = %s, prompt_name = %s WHERE user_id = %s AND client_id = %s AND prompt_name = %s AND template_name IS NULL",
                    (Json({'prompt': prompt_content}), prompt_name, current_user.id, client_id, original_prompt_name)
                )
                if cur.rowcount == 0:
                    flash(f'Prompt "{original_prompt_name}" not found for update', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('create_prompt', client_id=client_id))
                conn.commit()
                cur.close()
                conn.close()
                flash(f'Prompt "{prompt_name}" updated successfully', 'success')
                return redirect(url_for('create_prompt', client_id=client_id))
            except Exception as e:
                logger.error(f"Error updating prompt: {str(e)}")
                flash(f'Failed to update prompt: {str(e)}', 'danger')
                return redirect(url_for('create_prompt', client_id=client_id))
    
    conn = get_db_connection()
    cur = conn.cursor()
    if selected_client:
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt IS NOT NULL AND template_name IS NULL",
            (current_user.id, selected_client)
        )
    else:
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND client_id = '' AND prompt IS NOT NULL AND template_name IS NULL",
            (current_user.id,)
        )
    prompts = [{'prompt_name': row[0], 'prompt_content': row[1] or ''} for row in cur.fetchall()]
    cur.close()
    conn.close()
    logger.info(f"Prompts for client {selected_client or 'global'}: {prompts}")
    
    return render_template('create_prompt.html', clients=clients, selected_client=selected_client, prompts=prompts, selected_prompt=edit_prompt)

# Template creation/editing route
@app.route('/create_template', methods=['GET', 'POST'])
@login_required
def create_template():
    clients = get_user_clients(current_user.id)
    selected_client = request.args.get('client_id', '')
    edit_template = request.args.get('edit_template', '')
    templates = []
    prompts = []
    
    conn = get_db_connection()
    cur = conn.cursor()
    if selected_client:
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt IS NOT NULL AND template_name IS NULL",
            (current_user.id, selected_client)
        )
    else:
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND client_id = '' AND prompt IS NOT NULL AND template_name IS NULL",
            (current_user.id,)
        )
    prompts = [{'prompt_name': row[0], 'prompt_content': row[1] or ''} for row in cur.fetchall()]
    
    if request.method == 'POST':
        action = request.form.get('action')
        client_id = request.form.get('client_id', '').strip()
        template_name = request.form.get('template_name', '').strip()
        prompt_name = request.form.get('prompt_name', '').strip()
        prompt_name_new = request.form.get('prompt_name_new', '').strip()
        prompt_content = request.form.get('prompt_content', '').strip()
        template_file = request.files.get('template_file')
        original_template_name = request.form.get('original_template_name', template_name).strip()
        
        if prompt_name_new and prompt_content:
            prompt_name = prompt_name_new
        
        if action == 'create':
            if not template_name:
                flash('Template name cannot be empty', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
            if not template_file or not template_file.filename.endswith('.docx'):
                flash('Valid .docx template file required', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
            if not prompt_name:
                flash('Please select an existing prompt or provide a new one', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND template_name = %s",
                    (current_user.id, client_id, template_name)
                )
                if cur.fetchone():
                    flash(f'Template "{template_name}" already exists for this client', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('create_template', client_id=client_id))
                if prompt_name_new:
                    cur.execute(
                        "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND prompt_name = %s AND template_name IS NULL",
                        (current_user.id, client_id, prompt_name)
                    )
                    if cur.fetchone():
                        flash(f'Prompt "{prompt_name}" already exists for this client', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('create_template', client_id=client_id))
                    save_prompt(prompt_content, client_id, current_user.id, prompt_name)
                prompt_json = None
                if prompt_name and not prompt_name_new:
                    cur.execute(
                        "SELECT prompt FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt_name = %s AND prompt IS NOT NULL AND template_name IS NULL LIMIT 1",
                        (current_user.id, client_id, prompt_name)
                    )
                    result = cur.fetchone()
                    if result:
                        prompt_json = result[0]
                    else:
                        flash(f'Prompt "{prompt_name}" not found', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('create_template', client_id=client_id))
                file_data = template_file.read()
                cur.execute(
                    "INSERT INTO settings (user_id, client_id, prompt, prompt_name, template, template_name) "
                    "VALUES (%s, %s, %s, %s, %s, %s)",
                    (current_user.id, client_id, Json(prompt_json) if prompt_json else None, prompt_name, file_data, template_name)
                )
                conn.commit()
                # Clear cached styles
                cache_key = md5(f"{current_user.id}_{client_id}_{template_name}".encode()).hexdigest()
                cache_path = os.path.join(app.config['UPLOAD_FOLDER'], f'template_styles_{cache_key}.json')
                if os.path.exists(cache_path):
                    os.remove(cache_path)
                flash(f'Template "{template_name}" created successfully', 'success')
                cur.close()
                conn.close()
                return redirect(url_for('create_template', client_id=client_id))
            except Exception as e:
                logger.error(f"Error creating template: {str(e)}")
                flash(f'Failed to create template: {str(e)}', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
        
        elif action == 'update':
            if not template_name:
                flash('Template name cannot be empty', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
            if not prompt_name:
                flash('Please select an existing prompt or provide a new one', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                if template_name != original_template_name:
                    cur.execute(
                        "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND template_name = %s",
                        (current_user.id, client_id, template_name)
                    )
                    if cur.fetchone():
                        flash(f'Template "{template_name}" already exists for this client', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('create_template', client_id=client_id))
                if prompt_name_new:
                    cur.execute(
                        "SELECT id FROM settings WHERE user_id = %s AND client_id = %s AND prompt_name = %s AND template_name IS NULL",
                        (current_user.id, client_id, prompt_name)
                    )
                    if cur.fetchone():
                        flash(f'Prompt "{prompt_name}" already exists for this client', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('create_template', client_id=client_id))
                    save_prompt(prompt_content, client_id, current_user.id, prompt_name)
                prompt_json = None
                if prompt_name and not prompt_name_new:
                    cur.execute(
                        "SELECT prompt FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt_name = %s AND prompt IS NOT NULL AND template_name IS NULL LIMIT 1",
                        (current_user.id, client_id, prompt_name)
                    )
                    result = cur.fetchone()
                    if result:
                        prompt_json = result[0]
                if template_file and template_file.filename.endswith('.docx'):
                    cur.execute(
                        "UPDATE settings SET template = %s, prompt = %s, prompt_name = %s, template_name = %s "
                        "WHERE user_id = %s AND client_id = %s AND template_name = %s",
                        (template_file.read(), Json(prompt_json) if prompt_json else None, prompt_name, template_name, current_user.id, client_id, original_template_name)
                    )
                    # Clear cached styles
                    cache_key = md5(f"{current_user.id}_{client_id}_{template_name}".encode()).hexdigest()
                    cache_path = os.path.join(app.config['UPLOAD_FOLDER'], f'template_styles_{cache_key}.json')
                    if os.path.exists(cache_path):
                        os.remove(cache_path)
                else:
                    cur.execute(
                        "UPDATE settings SET prompt = %s, prompt_name = %s, template_name = %s "
                        "WHERE user_id = %s AND client_id = %s AND template_name = %s",
                        (Json(prompt_json) if prompt_json else None, prompt_name, template_name, current_user.id, client_id, original_template_name)
                    )
                if cur.rowcount == 0:
                    flash(f'Template "{original_template_name}" not found for update', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('create_template', client_id=client_id))
                conn.commit()
                cur.close()
                conn.close()
                flash(f'Template "{template_name}" updated successfully', 'success')
                return redirect(url_for('create_template', client_id=client_id))
            except Exception as e:
                logger.error(f"Error updating template: {str(e)}")
                flash(f'Failed to update template: {str(e)}', 'danger')
                return redirect(url_for('create_template', client_id=client_id))
    
    if selected_client:
        cur.execute(
            "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template IS NOT NULL",
            (current_user.id, selected_client)
        )
    else:
        cur.execute(
            "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND client_id = '' AND template IS NOT NULL",
            (current_user.id,)
        )
    templates = [{'template_name': row[0], 'prompt_name': row[1]} for row in cur.fetchall()]
    cur.close()
    conn.close()
    logger.info(f"Templates for client {selected_client or 'global'}: {templates}")
    
    return render_template('create_template.html', clients=clients, selected_client=selected_client, templates=templates, prompts=prompts, selected_template=edit_template)

# Template deletion route
@app.route('/delete_template', methods=['POST'])
@login_required
def delete_template():
    try:
        client_id = request.form.get('client_id', '').strip()
        template_name = request.form.get('template_name').strip()
        if not template_name:
            return jsonify({'success': False, 'error': 'Template name is required'}), 400
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM settings WHERE user_id = %s AND client_id = %s AND template_name = %s AND template IS NOT NULL",
            (current_user.id, client_id, template_name)
        )
        if cur.rowcount == 0:
            cur.close()
            conn.close()
            return jsonify({'success': False, 'error': 'Template not found'}), 404
        conn.commit()
        cur.close()
        conn.close()
        # Remove cached styles
        cache_key = md5(f"{current_user.id}_{client_id}_{template_name}".encode()).hexdigest()
        cache_path = os.path.join(app.config['UPLOAD_FOLDER'], f'template_styles_{cache_key}.json')
        if os.path.exists(cache_path):
            os.remove(cache_path)
        logger.info(f"Deleted template '{template_name}' for client {client_id or 'global'}, user {current_user.id}")
        return jsonify({'success': True}), 200
    except Exception as e:
        logger.error(f"Error deleting template: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Prompt deletion route
@app.route('/delete_prompt', methods=['POST'])
@login_required
def delete_prompt():
    try:
        client_id = request.form.get('client_id', '').strip()
        prompt_name = request.form.get('prompt_name').strip()
        if not prompt_name:
            return jsonify({'success': False, 'error': 'Prompt name is required'}), 400
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM settings WHERE user_id = %s AND client_id = %s AND prompt_name = %s "
            "AND prompt IS NOT NULL AND template_name IS NULL AND template IS NULL",
            (current_user.id, client_id, prompt_name)
        )
        if cur.rowcount == 0:
            cur.close()
            conn.close()
            return jsonify({'success': False, 'error': 'Prompt not found'}), 404
        conn.commit()
        cur.close()
        conn.close()
        logger.info(f"Deleted prompt '{prompt_name}' for client {client_id or 'global'}, user {current_user.id}")
        return jsonify({'success': True}), 200
    except Exception as e:
        logger.error(f"Error deleting prompt: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Client deletion route
@app.route('/delete_client', methods=['POST'])
@login_required
def delete_client():
    try:
        client_id = request.form.get('client_id').strip()
        if not client_id:
            return jsonify({'success': False, 'error': 'Client ID is required'}), 400
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM settings WHERE user_id = %s AND client_id = %s",
            (current_user.id, client_id)
        )
        if cur.rowcount == 0:
            cur.close()
            conn.close()
            return jsonify({'success': False, 'error': 'Client not found'}), 404
        conn.commit()
        cur.close()
        conn.close()
        # Remove cached styles for client
        cache_dir = app.config['UPLOAD_FOLDER']
        for f in os.listdir(cache_dir):
            if f.startswith(f'template_styles_{current_user.id}_{client_id}_'):
                os.remove(os.path.join(cache_dir, f))
        logger.info(f"Deleted client '{client_id}' for user {current_user.id}")
        return jsonify({'success': True}), 200
    except Exception as e:
        logger.error(f"Error deleting client: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Prompt management routes
@app.route('/load_prompts', methods=['POST'])
@login_required
def load_prompts():
    try:
        client_id = request.form.get('client_id', '')
        conn = get_db_connection()
        cur = conn.cursor()
        if client_id:
            cur.execute(
                "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt IS NOT NULL AND template_name IS NULL",
                (current_user.id, client_id)
            )
        else:
            cur.execute(
                "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND client_id = '' AND prompt IS NOT NULL AND template_name IS NULL",
                (current_user.id,)
            )
        prompts = [{'prompt_name': row[0], 'prompt_content': row[1] or ''} for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Loaded prompts for client {client_id or 'global'}: {prompts}")
        return jsonify({'prompts': prompts}), 200
    except Exception as e:
        logger.error(f"Error loading prompts: {str(e)}")
        return jsonify({'error': 'Failed to load prompts'}), 500

@app.route('/load_templates', methods=['POST'])
@login_required
def load_templates():
    try:
        client_id = request.form.get('client_id', '')
        conn = get_db_connection()
        cur = conn.cursor()
        if client_id:
            cur.execute(
                "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template IS NOT NULL",
                (current_user.id, client_id)
            )
        else:
            cur.execute(
                "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND client_id = '' AND template IS NOT NULL",
                (current_user.id,)
            )
        templates = [{'template_name': row[0], 'prompt_name': row[1]} for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Loaded templates for client {client_id or 'global'}: {templates}")
        return jsonify({'templates': templates}), 200
    except Exception as e:
        logger.error(f"Error loading templates: {str(e)}")
        return jsonify({'error': 'Failed to load templates'}), 500

@app.route('/load_client', methods=['POST'])
@login_required
def load_client():
    try:
        data = request.form
        client_id = data.get('client_id', '')
        template_name = data.get('template_name')
        prompt_name = data.get('prompt_name')
        if template_name:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content, prompt_name, template IS NOT NULL AS has_file "
                "FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template_name = %s LIMIT 1",
                (current_user.id, client_id, template_name)
            )
            result = cur.fetchone()
            cur.close()
            conn.close()
            if result:
                prompt = result[0] or ''
                prompt_name = result[1] or 'Custom'
                has_file = result[2]
                return jsonify({'prompt': prompt, 'prompt_name': prompt_name, 'has_file': has_file}), 200
            return jsonify({'prompt': '', 'prompt_name': 'Custom', 'has_file': False}), 404
        elif prompt_name and prompt_name != 'Custom':
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content FROM settings "
                "WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt_name = %s AND prompt IS NOT NULL AND template_name IS NULL "
                "ORDER BY client_id DESC LIMIT 1",
                (current_user.id, client_id, prompt_name)
            )
            result = cur.fetchone()
            cur.close()
            conn.close()
            if result:
                return jsonify({'prompt': result[0] or '', 'prompt_name': prompt_name, 'has_file': False}), 200
            return jsonify({'prompt': '', 'prompt_name': 'Custom', 'has_file': False}), 404
        return jsonify({'prompt': '', 'prompt_name': 'Custom', 'has_file': False}), 200
    except Exception as e:
        logger.error(f"Error loading client: {str(e)}")
        return jsonify({'error': 'Failed to load client'}), 500

# Main page route
@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    clients = get_user_clients(current_user.id)
    selected_client = request.form.get('client', '') if request.method == 'POST' else ''
    selected_template = None
    prompt_name = 'Custom'
    prompt_content = ''

    # Always fetch templates and prompts, including global ones
    templates = get_templates_for_client(selected_client, current_user.id)
    prompts = get_prompts_for_client(selected_client, current_user.id)

    if request.method == 'POST':
        selected_template = request.form.get('template')
        prompt_name = request.form.get('prompt_name', 'Custom')
        source_file = request.files.get('source_file')
        ai_prompt = request.form.get('ai_prompt')
        source_text = request.form.get('source_text', '')

        if source_file and source_file.filename.endswith('.docx'):
            template_name = selected_template if selected_template else None
            if template_name and not ai_prompt:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT prompt->'prompt' AS prompt_content "
                    "FROM settings "
                    "WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template_name = %s AND prompt IS NOT NULL LIMIT 1",
                    (current_user.id, selected_client, template_name)
                )
                result = cur.fetchone()
                cur.close()
                conn.close()
                if result:
                    ai_prompt = result[0]
                else:
                    flash('No prompt found for the selected template.', 'error')
                    return redirect(url_for('index'))

            try:
                content = process_docx(source_file)
                with TemporaryDirectory() as temp_dir:
                    temp_template_path = os.path.join(temp_dir, 'temp_template.docx')
                    fetch_template(temp_template_path, selected_client, current_user.id, template_name)
                    output_path = os.path.join(temp_dir, 'reformatted_document.docx')
                    sections = call_ai_api(content, selected_client, current_user.id, prompt_name, custom_prompt=ai_prompt, template_path=temp_template_path)
                    create_reformatted_docx(sections, output_path, selected_client, current_user.id, template_path=temp_template_path)
                    return send_file(output_path, as_attachment=True, download_name='reformatted_document.docx')
            except Exception as e:
                logger.error(f"Error processing document: {str(e)}")
                flash(f"Error processing document: {str(e)}", 'error')

        elif source_text:
            template_name = selected_template if selected_template else None
            if template_name and not ai_prompt:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "SELECT prompt->'prompt' AS prompt_content "
                    "FROM settings "
                    "WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template_name = %s AND prompt IS NOT NULL LIMIT 1",
                    (current_user.id, selected_client, template_name)
                )
                result = cur.fetchone()
                cur.close()
                conn.close()
                if result:
                    ai_prompt = result[0]
                else:
                    flash('No prompt found for the selected template.', 'error')
                    return redirect(url_for('index'))

            try:
                content = process_text_input(source_text)
                with TemporaryDirectory() as temp_dir:
                    temp_template_path = os.path.join(temp_dir, 'temp_template.docx')
                    fetch_template(temp_template_path, selected_client, current_user.id, template_name)
                    output_path = os.path.join(temp_dir, 'reformatted_document.docx')
                    sections = call_ai_api(content, selected_client, current_user.id, prompt_name, custom_prompt=ai_prompt, template_path=temp_template_path)
                    create_reformatted_docx(sections, output_path, selected_client, current_user.id, template_path=temp_template_path)
                    return send_file(output_path, as_attachment=True, download_name='reformatted_document.docx')
            except Exception as e:
                logger.error(f"Error processing text input: {str(e)}")
                flash(f"Error processing text input: {str(e)}", 'error')

        if selected_template:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content, prompt_name "
                "FROM settings "
                "WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template_name = %s LIMIT 1",
                (current_user.id, selected_client, selected_template)
            )
            result = cur.fetchone()
            cur.close()
            conn.close()
            if result:
                prompt_content = result[0] or ''
                prompt_name = result[1] or 'Custom'

    return render_template('index.html', clients=clients, templates=templates, prompts=prompts,
                         selected_client=selected_client, selected_template=selected_template,
                         prompt_name=prompt_name, prompt_content=prompt_content)

# API and document processing
AI_API_URL = os.environ.get('AI_API_URL', 'https://api.openai.com/v1/chat/completions')
API_KEY = os.environ.get('API_KEY', 'your-api-key')

DEFAULT_AI_PROMPT = """You are a document analyst. Analyze the provided document content and categorize it into sections based on the input text and tables. Return a JSON object with keys representing section headers and values containing the extracted or rewritten content. Preserve references separately. Ensure the response is valid JSON. For tables, return a dictionary where keys are descriptive section names and values are lists of rows, each row being a list of cell values. Focus on accurately interpreting and summarizing the source material."""

def load_prompt(client_id=None, user_id=None, prompt_name=None):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        if client_id is not None and user_id and prompt_name:
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content FROM settings WHERE client_id = %s AND user_id = %s AND prompt_name = %s AND prompt IS NOT NULL AND template_name IS NULL ORDER BY created_at DESC LIMIT 1",
                (client_id, user_id, prompt_name)
            )
        elif client_id is not None and user_id:
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content FROM settings WHERE client_id = %s AND user_id = %s AND prompt IS NOT NULL AND template_name IS NULL ORDER BY created_at DESC LIMIT 1",
                (client_id, user_id)
            )
        else:
            cur.execute(
                "SELECT prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND client_id = '' AND prompt IS NOT NULL AND template_name IS NULL ORDER BY created_at DESC LIMIT 1",
                (user_id,)
            )
        result = cur.fetchone()
        cur.close()
        conn.close()
        return result[0] if result and result[0] else DEFAULT_AI_PROMPT
    except Exception as e:
        logger.error(f"Error loading prompt: {str(e)}")
        return DEFAULT_AI_PROMPT

def save_prompt(prompt, client_id, user_id, prompt_name):
    try:
        if not prompt_name:
            raise ValueError(f"Prompt name cannot be empty: prompt_name={prompt_name}")
        if not prompt:
            raise ValueError(f"Prompt content cannot be empty")
        sanitized_prompt = prompt.encode('utf-8', errors='replace').decode('utf-8')
        client_id = client_id or ''
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO settings (user_id, client_id, prompt, prompt_name) VALUES (%s, %s, %s, %s) "
            "ON CONFLICT ON CONSTRAINT settings_unique_user_client_prompt_template DO UPDATE SET prompt = EXCLUDED.prompt",
            (user_id, client_id, Json({'prompt': sanitized_prompt}), prompt_name)
        )
        conn.commit()
        cur.close()
        conn.close()
        logger.info(f"Saved prompt '{prompt_name}' for client {client_id or 'global'}, user {user_id}: {sanitized_prompt[:100]}...")
    except Exception as e:
        logger.error(f"Error saving prompt: {str(e)}")
        raise

def get_user_clients(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT DISTINCT client_id FROM settings WHERE user_id = %s AND client_id IS NOT NULL AND client_id != ''",
            (user_id,)
        )
        clients = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Fetched clients for user {user_id}: {clients}")
        return clients
    except Exception as e:
        logger.error(f"Error getting clients: {str(e)}")
        return []

def get_templates_for_client(client_id, user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT template_name, prompt_name FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND template IS NOT NULL",
            (user_id, client_id or '')
        )
        templates = [{'template_name': row[0], 'prompt_name': row[1]} for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Fetched templates for client {client_id or 'global'}, user {user_id}: {templates}")
        return templates
    except Exception as e:
        logger.error(f"Error getting templates for client: {str(e)}")
        return []

def get_prompts_for_client(client_id, user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT prompt_name, prompt->'prompt' AS prompt_content FROM settings WHERE user_id = %s AND (client_id = %s OR client_id = '') AND prompt IS NOT NULL AND template_name IS NULL",
            (user_id, client_id or '')
        )
        prompts = [{'prompt_name': row[0], 'prompt_content': row[1] or ''} for row in cur.fetchall()]
        cur.close()
        conn.close()
        logger.info(f"Fetched prompts for client {client_id or 'global'}, user {user_id}: {prompts}")
        return prompts
    except Exception as e:
        logger.error(f"Error getting prompts for client: {str(e)}")
        return []

def process_text_input(text):
    try:
        doc = Document()
        paragraphs = text.split('\n')
        current_section = None
        content = {"text_chunks": [], "tables": [], "references": [], "section_order": []}
        chunk = []
        chunk_size = 1000
        current_chunk_length = 0

        for para in paragraphs:
            text = para.strip()
            if text:
                if text.lower().startswith("references"):
                    current_section = "References"
                    content["section_order"].append(current_section)
                    continue
                if current_section == "References":
                    content["references"].append(text)
                else:
                    # Simple heuristic for section detection in plain text
                    is_header = text.isupper() or len(text.split()) < 5
                    if is_header:
                        if current_section:
                            content["section_order"].append(current_section)
                        current_section = text
                    chunk.append(f"[{current_section}] {text}" if current_section else text)
                    current_chunk_length += len(text)
                    if current_chunk_length >= chunk_size and current_section:
                        content["text_chunks"].append("\n".join(chunk))
                        chunk = []
                        current_chunk_length = 0
        if chunk:
            content["text_chunks"].append("\n".join(chunk))
            if current_section and current_section not in content["section_order"]:
                content["section_order"].append(current_section)
        
        logger.info(f"Processed text input section order: {content['section_order']}")
        return content
    except Exception as e:
        logger.error(f"Error processing text input: {str(e)}")
        raise

def process_docx(file):
    try:
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        doc = Document(input_path)
        content = {"text_chunks": [], "tables": [], "references": [], "section_order": []}
        in_references = False
        chunk = []
        chunk_size = 1000
        current_chunk_length = 0
        current_section = None

        known_headers = ["introduction", "summary", "experience", "education", "affiliations", "skills", "competencies", "results", "conclusion"]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if text.lower().startswith("references"):
                    in_references = True
                    current_section = "References"
                    content["section_order"].append(current_section)
                    continue
                if in_references:
                    content["references"].append(text)
                else:
                    is_header = (
                        para.runs and (para.runs[0].bold or para.runs[0].font.size > Pt(12)) or
                        text.isupper() or
                        any(text.lower().startswith(h) for h in known_headers)
                    )
                    if is_header:
                        if current_section:
                            content["section_order"].append(current_section)
                        current_section = text
                    chunk.append(f"[{current_section}] {text}" if current_section else text)
                    current_chunk_length += len(text)
                    if current_chunk_length >= chunk_size and current_section:
                        content["text_chunks"].append("\n".join(chunk))
                        chunk = []
                        current_chunk_length = 0
        if chunk:
            content["text_chunks"].append("\n".join(chunk))
            if current_section and current_section not in content["section_order"]:
                content["section_order"].append(current_section)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_data.append(row_data)
            if table_data:
                logger.info(f"Extracted table: {table_data}")
                content["tables"].append(table_data)
        os.remove(input_path)
        logger.info(f"Source section order: {content['section_order']}")
        return content
    except Exception as e:
        logger.error(f"Error processing docx: {str(e)}")
        raise

def fetch_template(output_path, client_id=None, user_id=None, template_name=None):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        if client_id is not None and user_id and template_name:
            cur.execute(
                "SELECT template FROM settings WHERE client_id = %s AND user_id = %s AND template_name = %s AND template IS NOT NULL ORDER BY created_at DESC LIMIT 1",
                (client_id, user_id, template_name)
            )
        elif client_id is not None and user_id:
            cur.execute(
                "SELECT template FROM settings WHERE client_id = %s AND user_id = %s AND template IS NOT NULL ORDER BY created_at DESC LIMIT 1",
                (client_id, user_id)
            )
        else:
            cur.execute(
                "SELECT template FROM settings WHERE user_id = %s AND client_id = '' AND template IS NOT NULL ORDER BY created_at DESC LIMIT 1",
                (user_id,)
            )
        result = cur.fetchone()
        cur.close()
        conn.close()
        if result and result[0]:
            with open(output_path, 'wb') as f:
                f.write(result[0])
            logger.info(f"Successfully fetched template: {output_path}")
            return True
        logger.warning("No template found in database")
        return False
    except Exception as e:
        logger.error(f"Error fetching template: {str(e)}")
        return False

def extract_template_styles(template_path, user_id, client_id, template_name):
    cache_key = md5(f"{user_id}_{client_id}_{template_name}".encode()).hexdigest()
    cache_path = os.path.join(app.config['UPLOAD_FOLDER'], f'template_styles_{cache_key}.json')
    
    if os.path.exists(cache_path):
        with open(cache_path, 'r') as f:
            styles = json.load(f)
        logger.info(f"Loaded cached template styles: {cache_path}")
        return styles
    
    styles = {
        "sections": {},
        "default": {
            "font": "Arial",
            "size": Pt(11),
            "bold": False,
            "color": RGBColor(0, 0, 0),
            "spacing_before": Pt(6),
            "spacing_after": Pt(6),
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "is_horizontal_list": False,
            "is_table": False
        }
    }
    
    try:
        doc = Document(template_path) if os.path.exists(template_path) else Document()
        for para in doc.paragraphs:
            text = para.text.strip().lower()
            if not text or not para.runs:
                continue
            run = para.runs[0]
            section_name = None
            if para.runs and (run.bold or run.font.size > Pt(12) or text.isupper()):
                section_name = para.text.strip()
            style = {
                "font": run.font.name or "Arial",
                "size": run.font.size or Pt(11),
                "bold": run.bold if run.bold is not None else False,
                "color": run.font.color.rgb or RGBColor(0, 0, 0),
                "spacing_before": para.paragraph_format.space_before or Pt(6),
                "spacing_after": para.paragraph_format.space_after or Pt(6),
                "alignment": para.paragraph_format.alignment or WD_ALIGN_PARAGRAPH.LEFT,
                "is_horizontal_list": "•" in para.text and text.count('\n') <= 1,
                "is_table": False
            }
            if section_name:
                styles["sections"][section_name] = style
            elif "•" in para.text:
                styles["sections"]["list_item"] = style
        
        for table in doc.tables:
            if table.rows:
                cell = table.rows[0].cells[0]
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                style = {
                    "font": run.font.name or "Arial" if run else "Arial",
                    "size": run.font.size or Pt(11) if run else Pt(11),
                    "bold": run.bold if run and run.bold is not None else False,
                    "color": run.font.color.rgb or RGBColor(0, 0, 0) if run else RGBColor(0, 0, 0),
                    "spacing_before": Pt(6),
                    "spacing_after": Pt(6),
                    "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                    "is_horizontal_list": False,
                    "is_table": True
                }
                styles["sections"]["table"] = style
        
        with open(cache_path, 'w') as f:
            json.dump(styles, f, default=str)
        logger.info(f"Cached template styles: {styles}")
        return styles
    except Exception as e:
        logger.error(f"Error extracting template styles: {str(e)}")
        return styles

def generate_supplemental_prompt(template_structure, source_sections):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    prompt = f"""
You are an AI assistant tasked with generating a supplemental prompt to enhance document reformatting. Given the template structure and source sections below, create a concise prompt that:

1. Maps source sections to template sections semantically, based on their content and purpose.
2. Preserves the source's logical order of content.
3. Ensures all source content is included, creating new sections if needed to match the template's style.
4. Specifies formatting for sections (e.g., lists, tables) based on the template's structure.
5. Avoids including font/size/color details, as these are handled by the application.

**Template Structure**: {json.dumps(template_structure, indent=2)}
**Source Sections**: {json.dumps(source_sections, indent=2)}

Output a plain text prompt, starting with "Supplemental Instructions:".
"""
    payload = {
        "model": "gpt-4o",
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": "Generate the supplemental prompt."}
        ],
        "max_tokens": 500,
        "temperature": 0.7
    }
    try:
        response = requests.post(AI_API_URL, headers=headers, json=payload, timeout=10)
        response.raise_for_status()
        data = response.json()
        if "choices" in data and data["choices"]:
            supplemental = data["choices"][0]["message"]["content"]
            logger.info(f"Generated supplemental prompt: {supplemental[:200]}...")
            return supplemental
        logger.error("No choices in supplemental prompt response")
        return ""
    except Exception as e:
        logger.error(f"Error generating supplemental prompt: {str(e)}")
        return ""

def call_ai_api(content, client_id=None, user_id=None, prompt_name=None, custom_prompt=None, template_path=None):
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    base_prompt = custom_prompt if custom_prompt else load_prompt(client_id, user_id, prompt_name)
    
    # Extract template structure
    template_doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
    template_structure = {"sections": []}
    for para in template_doc.paragraphs:
        text = para.text.strip()
        if text and (para.runs and (para.runs[0].bold or para.runs[0].font.size > Pt(12)) or text.isupper()):
            template_structure["sections"].append({
                "name": text,
                "is_list": "•" in para.text,
                "is_table": False
            })
    for table in template_doc.tables:
        if table.rows:
            template_structure["sections"].append({
                "name": "Table",
                "is_list": False,
                "is_table": True
            })
    
    # Prepare source sections
    source_sections = content.get("section_order", [])
    
    # Generate supplemental prompt
    supplemental_prompt = generate_supplemental_prompt(template_structure, source_sections)
    full_prompt = f"{base_prompt}\n\n{supplemental_prompt}"
    
    def process_chunk(chunk_text, tables, chunk_index):
        messages = [
            {
                "role": "system",
                "content": full_prompt
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": f"Input Text (Chunk {chunk_index}):\n{chunk_text}\n\nTables:\n{json.dumps(tables)}"}
                ]
            }
        ]
        payload = {
            "model": "gpt-4o",
            "messages": messages,
            "max_tokens": 4096,
            "temperature": 0.7,
            "response_format": {"type": "json_object"}
        }
        try:
            response = requests.post(AI_API_URL, headers=headers, json=payload, timeout=15)
            response.raise_for_status()
            if not response.text:
                logger.error(f"Empty response from API (chunk {chunk_index})")
                return {"error": f"Empty response from API in chunk {chunk_index}"}
            try:
                data = response.json()
                if "choices" not in data or not data["choices"]:
                    logger.error(f"No choices in API response (chunk {chunk_index}): {data}")
                    return {"error": f"No valid choices in API response for chunk {chunk_index}"}
                raw_content = data["choices"][0]["message"]["content"]
                logger.info(f"Raw AI response (chunk {chunk_index}): {raw_content[:1000]}...")
                try:
                    parsed_content = json.loads(raw_content)
                    if not isinstance(parsed_content, dict):
                        raise ValueError("AI response is not a JSON object")
                    normalized_content = {k.lower(): v for k, v in parsed_content.items()}
                    return {"index": chunk_index, "content": normalized_content}
                except json.JSONDecodeError as e:
                    logger.error(f"JSON validation error (chunk {chunk_index}): {str(e)}, raw: {raw_content[:1000]}")
                    try:
                        cleaned_content = raw_content.strip()
                        if not cleaned_content.endswith('}'):
                            cleaned_content = cleaned_content.rsplit('}', 1)[0] + '}'
                        if not cleaned_content.startswith('{'):
                            cleaned_content = '{' + cleaned_content.lstrip('{')
                        parsed_content = json.loads(cleaned_content)
                        if not isinstance(parsed_content, dict):
                            raise ValueError("Cleaned AI response is not a JSON object")
                        normalized_content = {k.lower(): v for k, v in parsed_content.items()}
                        return {"index": chunk_index, "content": normalized_content}
                    except json.JSONDecodeError as e2:
                        logger.error(f"Failed to fix JSON (chunk {chunk_index}): {str(e2)}")
                        return {"error": f"Invalid JSON in chunk {chunk_index}: {str(e)}"}
            except ValueError as e:
                logger.error(f"Invalid JSON response (chunk {chunk_index}): {response.text[:1000]}")
                return {"error": f"Invalid JSON response in chunk {chunk_index}: {response.text[:1000]}"}
        except requests.exceptions.HTTPError as e:
            error_response = response.text
            logger.error(f"API Error (chunk {chunk_index}): {response.status_code} - {error_response}")
            return {"error": f"HTTP Error in chunk {chunk_index}: {response.status_code} - {error_response}"}
        except Exception as e:
            logger.error(f"Unexpected Error (chunk {chunk_index}): {str(e)}")
            return {"error": str(e)}

    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as executor:
        futures = [executor.submit(process_chunk, chunk_text, content["tables"], i) for i, chunk_text in enumerate(content["text_chunks"], 1)]
        for future in concurrent.futures.as_completed(futures):
            results.append(future.result())
    
    # Sort results by chunk index to preserve source order
    results.sort(key=lambda x: x.get("index", float('inf')))
    
    merged_content = {}
    errors = []
    processed_items = set()
    for result in results:
        if "error" in result:
            errors.append(result["error"])
        else:
            for key, value in result["content"].items():
                if key in merged_content:
                    if isinstance(value, str) and isinstance(merged_content[key], str):
                        if value not in processed_items:
                            merged_content[key] = value + "\n" + merged_content[key]
                            processed_items.add(value)
                    elif isinstance(value, dict) and isinstance(merged_content[key], dict):
                        for subkey, subvalue in value.items():
                            if isinstance(subvalue, str) and isinstance(merged_content[key].get(subkey, ""), str):
                                if subvalue not in processed_items:
                                    merged_content[key][subkey] = subvalue + "\n" + merged_content[key].get(subkey, "")
                                    processed_items.add(subvalue)
                            elif isinstance(subvalue, list) and isinstance(merged_content[key].get(subkey, []), list):
                                new_items = []
                                for item in subvalue:
                                    if isinstance(item, dict):
                                        item_str = json.dumps(item, sort_keys=True)
                                        if item_str not in processed_items:
                                            new_items.append(item)
                                            processed_items.add(item_str)
                                    elif item not in processed_items:
                                        new_items.append(item)
                                        processed_items.add(item)
                                merged_content[key][subkey] = new_items + merged_content[key].get(subkey, [])
                            else:
                                merged_content[key][subkey] = subvalue
                                if isinstance(subvalue, str):
                                    processed_items.add(subvalue)
                    elif isinstance(value, list) and isinstance(merged_content[key], list):
                        new_items = []
                        for item in value:
                            if isinstance(item, dict):
                                item_str = json.dumps(item, sort_keys=True)
                                if item_str not in processed_items:
                                    new_items.append(item)
                                    processed_items.add(item_str)
                            elif item not in processed_items:
                                new_items.append(item)
                                processed_items.add(item)
                        merged_content[key] = new_items + merged_content[key]
                    else:
                        merged_content[key] = value
                        if isinstance(value, str):
                            processed_items.add(value)
                else:
                    merged_content[key] = value
                    if isinstance(value, str):
                        processed_items.add(value)
    
    if errors:
        logger.error(f"Errors in processing: {errors}")
        return {
            "error": "; ".join(errors),
            "content": content["text_chunks"][0][:500] if content["text_chunks"] else "",
            "tables": content["tables"],
            "references": content["references"],
            "section_order": content["section_order"]
        }
    
    merged_content["references"] = content["references"]
    merged_content["section_order"] = content["section_order"]
    logger.info(f"Merged AI response: {merged_content}")
    return merged_content

def create_reformatted_docx(sections, output_path, client_id=None, user_id=None, template_path=None):
    try:
        template_doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        doc = Document()
        
        styles = extract_template_styles(template_path, user_id, client_id, sections.get("template_name", "unknown"))
        
        def apply_template_style(paragraph, section_name):
            style = styles["sections"].get(section_name.lower(), styles["default"])
            for run in paragraph.runs:
                run.bold = style["bold"]
                run.font.name = style["font"]
                run.font.size = style["size"]
                run.font.color.rgb = style["color"]
            paragraph.paragraph_format.space_before = style["spacing_before"]
            paragraph.paragraph_format.space_after = style["spacing_after"]
            paragraph.paragraph_format.alignment = style["alignment"]
            paragraph.style = "Normal"
        
        def format_content(content, indent=0, section_name=None):
            if isinstance(content, str):
                return content.strip()
            elif isinstance(content, list) and section_name in styles["sections"] and styles["sections"].get(section_name.lower(), {}).get("is_horizontal_list", False):
                return " • ".join(item.strip() for item in content if isinstance(item, str))
            elif isinstance(content, list):
                formatted = []
                for item in content:
                    if isinstance(item, str):
                        formatted.append("• " + item.strip())
                    elif isinstance(item, dict):
                        sub_formatted = []
                        for subkey, subvalue in item.items():
                            subcontent = format_content(subvalue, indent + 1, section_name)
                            if subcontent:
                                sub_formatted.append(f"{subkey}: {subcontent}")
                        formatted.append("\n".join(sub_formatted))
                    elif isinstance(item, list):
                        formatted.extend(format_content(item, indent + 1, section_name))
                return "\n".join(formatted)
            elif isinstance(content, dict):
                formatted = []
                for subkey, subvalue in content.items():
                    subcontent = format_content(subvalue, indent + 1, section_name)
                    if subcontent:
                        formatted.append(f"{'  ' * indent}{subkey}: {subcontent}")
                return "\n".join(formatted)
            return ""

        if "error" in sections:
            para = doc.add_paragraph(f"Error: {sections['error']}")
            apply_template_style(para, "default")
            doc.save(output_path)
            return
        
        processed_keys = set()
        section_order = sections.get("section_order", [])
        # Add any sections in the JSON not in section_order
        for key in sections.keys():
            if key not in section_order and key not in ["references", "section_order", "error", "content", "tables"]:
                section_order.append(key)
        
        for key in section_order:
            normalized_key = key.lower()
            if normalized_key not in sections or normalized_key in processed_keys:
                continue
            value = sections[normalized_key]
            if not value or (isinstance(value, str) and not value.strip()) or (isinstance(value, list) and not value):
                continue
            processed_keys.add(normalized_key)
            
            display_key = key.capitalize()
            
            if display_key.lower() in styles["sections"]:
                para = doc.add_paragraph(display_key)
                apply_template_style(para, display_key)
                formatted_content = format_content(value, section_name=display_key)
                if formatted_content:
                    if styles["sections"].get(display_key.lower(), {}).get("is_horizontal_list", False):
                        para = doc.add_paragraph(formatted_content)
                        apply_template_style(para, display_key)
                    elif styles["sections"].get(display_key.lower(), {}).get("is_table", False):
                        table_data = value if isinstance(value, list) else [value]
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 1)
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                table.rows[i].cells[j].text = str(cell_data)
                        apply_template_style(table.rows[0].cells[0].paragraphs[0], display_key)
                    else:
                        lines = formatted_content.split("\n")
                        for line in lines:
                            if line.strip():
                                para = doc.add_paragraph(line, style="List Bullet" if line.startswith("•") else None)
                                apply_template_style(para, display_key)
            else:
                para = doc.add_paragraph(display_key)
                apply_template_style(para, "default")
                formatted_content = format_content(value, section_name=display_key)
                if formatted_content:
                    lines = formatted_content.split("\n")
                    for line in lines:
                        if line.strip():
                            para = doc.add_paragraph(line, style="List Bullet" if line.startswith("•") else None)
                            apply_template_style(para, "default")
        
        doc.save(output_path)
        logger.info(f"Successfully created reformatted document: {output_path}")
    except Exception as e:
        logger.error(f"Error creating reformatted docx: {str(e)}")
        raise

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)