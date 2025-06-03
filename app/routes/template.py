from flask import Blueprint, render_template, request, redirect, url_for, flash, Response
from flask_login import login_required, current_user
from ..utils.database import get_db_connection, get_user_clients, get_templates_for_client
from ..utils.document import process_docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import json
import os
from tempfile import NamedTemporaryFile
from hashlib import md5
from io import BytesIO
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import secrets
import logging

logger = logging.getLogger(__name__)

template_bp = Blueprint('template', __name__)

@template_bp.route('/create_template', methods=['GET', 'POST'])
@login_required
def create_template():
    clients = get_user_clients(current_user.id)
    selected_client = request.args.get('client_id', '') if request.method == 'GET' else request.form.get('client_id', '')
    edit_template = request.args.get('edit_template', '')
    reset_form = request.args.get('reset_form', 'false') == 'true' or not edit_template
    templates = []
    prompts = []
    conversion_prompts = []

    conn = get_db_connection()
    cur = conn.cursor()
    if selected_client:
        cur.execute(
            "SELECT p.id, p.prompt_name, p.prompt_type, p.content "
            "FROM prompts p LEFT JOIN clients c ON p.client_id = c.id "
            "WHERE p.user_id = %s AND (c.client_id = %s OR p.client_id IS NULL) AND p.prompt_type = 'template'",
            (current_user.id, selected_client)
        )
    else:
        cur.execute(
            "SELECT p.id, p.prompt_name, p.prompt_type, p.content "
            "FROM prompts p "
            "WHERE p.user_id = %s AND p.client_id IS NULL AND p.prompt_type = 'template'",
            (current_user.id,)
        )
    prompts = [{'id': row[0], 'prompt_name': row[1], 'prompt_type': row[2], 'content': row[3]} for row in cur.fetchall()]

    if request.method == 'POST':
        action = request.form.get('action')
        client_id = request.form.get('client_id', '').strip()
        template_name = request.form.get('template_name', '').strip()
        template_prompt_id = request.form.get('template_prompt_id', '').strip()
        conversion_prompt_id = request.form.get('conversion_prompt_id', '').strip()
        template_file = request.files.get('template_file')
        original_template_name = request.form.get('original_template_name', template_name).strip()

        client_id_value = None
        if client_id:
            cur.execute("SELECT id FROM clients WHERE user_id = %s AND client_id = %s", (current_user.id, client_id))
            client = cur.fetchone()
            if client:
                client_id_value = client[0]
            else:
                flash(f'Client "{client_id}" not found', 'danger')
                cur.close()
                conn.close()
                return redirect(url_for('template.create_template', client_id=client_id))

        if action == 'create':
            if not template_name:
                flash('Template name cannot be empty', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))
            if not template_prompt_id:
                flash('Template prompt is required', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))
            try:
                cur.execute(
                    "SELECT id FROM templates WHERE user_id = %s AND (client_id = %s OR %s IS NULL AND client_id IS NULL) AND template_name = %s",
                    (current_user.id, client_id_value, client_id_value, template_name)
                )
                if cur.fetchone():
                    flash(f'Template "{template_name}" already exists', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('template.create_template', client_id=client_id))
                file_data = template_file.read() if template_file and template_file.filename.endswith('.docx') else None
                if file_data:
                    logger.info(f"Storing template file for '{template_name}' (size: {len(file_data)} bytes)")
                else:
                    logger.info(f"No template file provided for '{template_name}'")
                cur.execute(
                    "INSERT INTO templates (user_id, client_id, template_name, template_prompt_id, template_file) "
                    "VALUES (%s, %s, %s, %s, %s) RETURNING id",
                    (current_user.id, client_id_value, template_name, int(template_prompt_id), file_data)
                )
                template_id = cur.fetchone()[0]
                if conversion_prompt_id:
                    cur.execute(
                        "INSERT INTO template_prompt_associations (template_id, conversion_prompt_id) VALUES (%s, %s)",
                        (template_id, int(conversion_prompt_id))
                    )
                conn.commit()
                flash(f'Template "{template_name}" created successfully', 'success')
                cur.close()
                conn.close()
                return redirect(url_for('template.create_template', client_id=client_id))
            except Exception as e:
                flash(f'Failed to create template: {str(e)}', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))

        elif action == 'update':
            if not template_name:
                flash('Template name cannot be empty', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))
            if not template_prompt_id:
                flash('Template prompt is required', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))
            try:
                if template_name != original_template_name:
                    cur.execute(
                        "SELECT id FROM templates WHERE user_id = %s AND (client_id = %s OR %s IS NULL AND client_id IS NULL) AND template_name = %s",
                        (current_user.id, client_id_value, client_id_value, template_name)
                    )
                    if cur.fetchone():
                        flash(f'Template "{template_name}" already exists', 'danger')
                        cur.close()
                        conn.close()
                        return redirect(url_for('template.create_template', client_id=client_id))
                cur.execute(
                    "SELECT id FROM templates WHERE user_id = %s AND (client_id = %s OR %s IS NULL AND client_id IS NULL) AND template_name = %s",
                    (current_user.id, client_id_value, client_id_value, original_template_name)
                )
                template = cur.fetchone()
                if not template:
                    flash(f'Template "{original_template_name}" not found', 'danger')
                    cur.close()
                    conn.close()
                    return redirect(url_for('template.create_template', client_id=client_id))
                template_id = template[0]
                file_data = template_file.read() if template_file and template_file.filename.endswith('.docx') else None
                if file_data:
                    logger.info(f"Updating template file for '{template_name}' (size: {len(file_data)} bytes)")
                    cur.execute(
                        "UPDATE templates SET template_name = %s, template_prompt_id = %s, template_file = %s "
                        "WHERE id = %s",
                        (template_name, int(template_prompt_id), file_data, template_id)
                    )
                else:
                    logger.info(f"No new template file provided for '{template_name}' during update")
                    cur.execute(
                        "UPDATE templates SET template_name = %s, template_prompt_id = %s "
                        "WHERE id = %s",
                        (template_name, int(template_prompt_id), template_id)
                    )
                cur.execute(
                    "DELETE FROM template_prompt_associations WHERE template_id = %s",
                    (template_id,)
                )
                if conversion_prompt_id:
                    cur.execute(
                        "INSERT INTO template_prompt_associations (template_id, conversion_prompt_id) VALUES (%s, %s)",
                        (template_id, int(conversion_prompt_id))
                    )
                conn.commit()
                flash(f'Template "{template_name}" updated successfully', 'success')
                cur.close()
                conn.close()
                return redirect(url_for('template.create_template', client_id=client_id))
            except Exception as e:
                flash(f'Failed to update template: {str(e)}', 'danger')
                return redirect(url_for('template.create_template', client_id=client_id))

    templates = get_templates_for_client(selected_client, current_user.id)
    logger.debug(f"Templates fetched: {templates}")
    filtered_templates = [
        template for template in templates
        if (selected_client and ('client_id' in template and (template['client_id'] is None or template['client_id'] == selected_client))) or
           (not selected_client and ('client_id' in template and template['client_id'] is None))
    ]
    cur.close()
    conn.close()

    return render_template('create_template.html', clients=clients, selected_client=selected_client, templates=filtered_templates, prompts=prompts, selected_template=edit_template if not reset_form else None)

@template_bp.route('/create_template_file/<int:template_id>', methods=['POST'])
@login_required
def create_template_file(template_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT t.template_name, p.content "
            "FROM templates t JOIN prompts p ON t.template_prompt_id = p.id "
            "WHERE t.id = %s AND t.user_id = %s",
            (template_id, current_user.id)
        )
        template = cur.fetchone()
        if not template:
            flash('Template not found', 'danger')
            cur.close()
            conn.close()
            return redirect(url_for('template.create_template'))

        template_name, prompt_content = template
        doc = Document()
        headers = {
            "Authorization": f"Bearer {os.environ.get('API_KEY')}",
            "Content-Type": "application/json"
        }
        
        # Construct prompt for LLM to generate .docx structure
        system_prompt = (
            "You are an AI tasked with generating a JSON object describing a .docx file structure based on a template prompt. "
            "The structure should include sections, content placeholders, and styling (font, size, bold, color in RGB, alignment, spacing, etc.). "
            "Use the prompt to infer the document's layout, sections, and semantic understanding of what each section should contain.\n\n"
            "**Template Prompt**: " + prompt_content + "\n\n"
            "**Output Format**:\n"
            "```json\n"
            "{\n"
            "  \"sections\": [\n"
            "    {\n"
            "      \"header\": \"Section Name\",\n"
            "      \"content\": [\"Placeholder text or instructions\"],\n"
            "      \"style\": {\n"
            "        \"font\": \"Font Name\",\n"
            "        \"size_pt\": Number,\n"
            "        \"bold\": Boolean,\n"
            "        \"color_rgb\": [R, G, B],\n"
            "        \"alignment\": \"left|center|right|justify\",\n"
            "        \"spacing_before_pt\": Number,\n"
            "        \"spacing_after_pt\": Number,\n"
            "        \"is_horizontal_list\": Boolean\n"
            "      }\n"
            "    },\n"
            "    ...\n"
            "  ]\n"
            "}\n"
            "```\n"
        )
        
        payload = {
            "model": "gpt-4o",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": "Generate the .docx structure based on the template prompt."}
            ],
            "max_tokens": 1500,
            "temperature": 0.7
        }
        
        session = requests.Session()
        retries = Retry(total=3, backoff_factor=1, status_forcelist=[429, 500, 502, 503, 504])
        session.mount('https://', HTTPAdapter(max_retries=retries))
        response = session.post(os.environ.get('AI_API_URL', 'https://api.openai.com/v1/chat/completions'), headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        data = response.json()
        if "choices" not in data or not data["choices"]:
            raise ValueError("No response from AI")
        doc_structure = json.loads(data["choices"][0]["message"]["content"])

        # Generate .docx file using python-docx
        for section in doc_structure.get("sections", []):
            style = section.get("style", {})
            para = doc.add_paragraph(section["header"])
            run = para.runs[0]
            run.font.name = style.get("font", "Arial")
            run.font.size = Pt(style.get("size_pt", 12))
            run.bold = style.get("bold", False)
            run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
            para.paragraph_format.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
            para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
            para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

            content = section.get("content", [])
            if style.get("is_horizontal_list", False):
                para = doc.add_paragraph(" • ".join(content))
                run = para.runs[0]
                run.font.name = style.get("font", "Arial")
                run.font.size = Pt(style.get("size_pt", 11))
                run.bold = style.get("bold", False)
                run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                para.paragraph_format.alignment = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))
            else:
                for line in content:
                    para = doc.add_paragraph(line, style="List Bullet" if line.startswith("•") else None)
                    run = para.runs[0]
                    run.font.name = style.get("font", "Arial")
                    run.font.size = Pt(style.get("size_pt", 11))
                    run.bold = style.get("bold", False)
                    run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                    para.paragraph_format.alignment = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                    para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                    para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

        with NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            doc.save(temp_file.name)
            with open(temp_file.name, 'rb') as f:
                file_data = f.read()
            os.unlink(temp_file.name)

        cur.execute(
            "UPDATE templates SET template_file = %s WHERE id = %s AND user_id = %s",
            (file_data, template_id, current_user.id)
        )
        logger.info(f"Updated template file for template ID {template_id} (size: {len(file_data)} bytes)")
        conn.commit()
        flash('Template file created successfully from prompt', 'success')
        cur.close()
        conn.close()
        return redirect(url_for('template.create_template'))
    except Exception as e:
        flash(f'Failed to create template file: {str(e)}', 'danger')
        if 'cur' in locals():
            cur.close()
        if 'conn' in locals():
            conn.close()
        return redirect(url_for('template.create_template'))

@template_bp.route('/create_prompt_from_file/<int:template_id>', methods=['POST'])
@login_required
def create_prompt_from_file(template_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT template_file, template_name, client_id FROM templates WHERE id = %s AND user_id = %s",
            (template_id, current_user.id)
        )
        template = cur.fetchone()
        if not template or not template[0]:
            flash('Template file not found', 'danger')
            cur.close()
            conn.close()
            return redirect(url_for('template.create_template'))

        template_file, template_name, client_id = template
        logger.info(f"Retrieved template file for template ID {template_id} (size: {len(template_file)} bytes)")
        with NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(template_file)
            temp_file_path = temp_file.name

        doc = Document(temp_file_path)
        os.unlink(temp_file_path)
        sections = []
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_header = (
                para.runs and (
                    para.runs[0].bold or
                    (para.runs[0].font.size is not None and para.runs[0].font.size > Pt(12))
                ) or
                text.isupper()
            )
            if is_header:
                current_section = text
                style = {
                    "font": para.runs[0].font.name or "Arial" if para.runs else "Arial",
                    "size_pt": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 12,
                    "bold": para.runs[0].bold if para.runs and para.runs[0].bold is not None else False,
                    "color_rgb": [para.runs[0].font.color.rgb.red, para.runs[0].font.color.rgb.green, para.runs[0].font.color.rgb.blue] if para.runs and para.runs[0].font.color.rgb else [0, 0, 0],
                    "alignment": {WD_ALIGN_PARAGRAPH.LEFT: "left", WD_ALIGN_PARAGRAPH.CENTER: "center", WD_ALIGN_PARAGRAPH.RIGHT: "right", WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"}.get(para.paragraph_format.alignment, "left"),
                    "spacing_before_pt": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 6,
                    "spacing_after_pt": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 6,
                    "is_horizontal_list": "•" in text and text.count('\n') <= 1
                }
                sections.append({"header": current_section, "style": style, "content": []})
            elif current_section:
                sections[-1]["content"].append(text)

        prompt_content = (
            "This is a template prompt for generating a document with the following structure and styling:\n\n"
            "The document should have the following sections, each with specific styling and semantic purposes:\n\n"
        )
        for section in sections:
            prompt_content += f"**Section: {section['header']}**\n"
            prompt_content += f"- **Purpose**: This section represents {section['header'].lower().replace(' ', '_')} content (e.g., if the section is 'Professional Experience', it should contain job roles, responsibilities, achievements).\n"
            prompt_content += "- **Style**:\n"
            prompt_content += f"  - Font: {section['style']['font']}\n"
            prompt_content += f"  - Size: {section['style']['size_pt']}pt\n"
            prompt_content += f"  - Bold: {section['style']['bold']}\n"
            prompt_content += f"  - Color: RGB({section['style']['color_rgb'][0]}, {section['style']['color_rgb'][1]}, {section['style']['color_rgb'][2]})\n"
            prompt_content += f"  - Alignment: {section['style']['alignment']}\n"
            prompt_content += f"  - Spacing Before: {section['style']['spacing_before_pt']}pt\n"
            prompt_content += f"  - Spacing After: {section['style']['spacing_after_pt']}pt\n"
            prompt_content += f"  - Horizontal List: {section['style']['is_horizontal_list']}\n"
            prompt_content += f"- **Content Placeholder**: {', '.join(section['content']) if section['content'] else 'Placeholder for relevant content'}\n\n"

        prompt_name = f"{template_name}_auto_prompt"
        cur.execute(
            "SELECT id FROM prompts WHERE user_id = %s AND (client_id = %s OR %s IS NULL AND client_id IS NULL) AND prompt_name = %s AND prompt_type = 'template'",
            (current_user.id, client_id, client_id, prompt_name)
        )
        if cur.fetchone():
            prompt_name = f"{prompt_name}_{secrets.token_hex(4)}"
        cur.execute(
            "INSERT INTO prompts (user_id, client_id, prompt_name, prompt_type, content) VALUES (%s, %s, %s, 'template', %s) RETURNING id",
            (current_user.id, client_id, prompt_name, 'template', prompt_content)
        )
        new_prompt_id = cur.fetchone()[0]
        cur.execute(
            "UPDATE templates SET template_prompt_id = %s WHERE id = %s",
            (new_prompt_id, template_id)
        )
        conn.commit()
        flash(f'Template prompt "{prompt_name}" created successfully from file', 'success')
        cur.close()
        conn.close()
        return redirect(url_for('template.create_template'))
    except Exception as e:
        flash(f'Failed to create prompt: {str(e)}', 'danger')
        if 'cur' in locals():
            cur.close()
        if 'conn' in locals():
            conn.close()
        return redirect(url_for('template.create_template'))

@template_bp.route('/view_template_file/<int:template_id>')
@login_required
def view_template_file(template_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT template_file, template_name FROM templates WHERE id = %s AND user_id = %s",
            (template_id, current_user.id)
        )
        template = cur.fetchone()
        cur.close()
        conn.close()
        if not template or not template[0]:
            flash('Template file not found', 'danger')
            return redirect(url_for('template.create_template'))
        file_data, template_name = template
        logger.info(f"Retrieved template file for download, template ID {template_id} (size: {len(file_data)} bytes)")
        return Response(
            file_data,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={template_name}.docx'}
        )
    except Exception as e:
        flash(f'Failed to view template file: {str(e)}', 'danger')
        return redirect(url_for('template.create_template'))

@template_bp.route('/delete_template/<int:template_id>', methods=['POST'])
@login_required
def delete_template(template_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM templates WHERE user_id = %s AND id = %s",
            (current_user.id, template_id)
        )
        if cur.rowcount == 0:
            flash(f'Template not found', 'danger')
        else:
            flash(f'Template deleted successfully', 'success')
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        flash(f'Failed to delete template: {str(e)}', 'danger')
    return redirect(url_for('template.create_template'))