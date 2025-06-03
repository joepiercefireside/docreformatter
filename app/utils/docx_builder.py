from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

def create_reformatted_docx(converted_content, template_file):
    """
    Create a reformatted .docx file by applying styles from the template file to the converted content.
    
    Args:
        converted_content (dict): Structured content in JSON format with sections.
        template_file (bytes): The template .docx file as a byte string.
    
    Returns:
        bytes: The reformatted .docx file as a byte string.
    """
    try:
        # Load the template file
        template_stream = BytesIO(template_file)
        template_doc = Document(template_stream)

        # Extract styles from the template
        styles = {}
        for para in template_doc.paragraphs:
            if not para.text.strip():
                continue
            # Determine if this paragraph is a header (based on bold, size, or uppercase)
            is_header = (
                para.runs and (
                    para.runs[0].bold or
                    (para.runs[0].font.size and para.runs[0].font.size.pt > 12) or
                    para.text.isupper()
                )
            )
            style_type = "header" if is_header else "body"
            
            # Extract style properties
            run = para.runs[0] if para.runs else None
            style = {
                "font_name": run.font.name if run and run.font.name else "Arial",
                "font_size_pt": run.font.size.pt if run and run.font.size else 11,
                "bold": run.bold if run and run.bold is not None else False,
                "color_rgb": (
                    [run.font.color.rgb.red, run.font.color.rgb.green, run.font.color.rgb.blue]
                    if run and run.font.color and run.font.color.rgb
                    else [0, 0, 0]
                ),
                "alignment": {
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
                }.get(para.paragraph_format.alignment, "left"),
                "spacing_before_pt": (
                    para.paragraph_format.space_before.pt
                    if para.paragraph_format.space_before
                    else 6
                ),
                "spacing_after_pt": (
                    para.paragraph_format.space_after.pt
                    if para.paragraph_format.space_after
                    else 6
                ),
                "is_horizontal_list": "•" in para.text and para.text.count('\n') <= 1
            }
            styles[style_type] = style

        logger.debug(f"Extracted styles from template: {styles}")

        # Create a new document for the output
        doc = Document()

        # Apply styles to the converted content
        # First, add the name (header style, typically larger and centered)
        if "name" in converted_content["sections"]:
            para = doc.add_paragraph(converted_content["sections"]["name"])
            style = styles.get("header", {})
            run = para.runs[0]
            run.font.name = style.get("font_name", "Arial")
            run.font.size = Pt(style.get("font_size_pt", 14) + 2)  # Slightly larger for name
            run.bold = style.get("bold", True)
            run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
            para.paragraph_format.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }.get("center", WD_ALIGN_PARAGRAPH.CENTER)  # Name is typically centered
            para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
            para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

        # Add contact info (header style, centered)
        if "contact" in converted_content["sections"]:
            para = doc.add_paragraph(converted_content["sections"]["contact"])
            style = styles.get("header", {})
            run = para.runs[0]
            run.font.name = style.get("font_name", "Arial")
            run.font.size = Pt(style.get("font_size_pt", 11))
            run.bold = False  # Contact info typically not bold
            run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
            para.paragraph_format.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }.get("center", WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
            para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

        # Add sections (e.g., professional summary, core competencies, etc.)
        for section_key, section_content in converted_content["sections"].items():
            if section_key in ["name", "contact"]:
                continue  # Already handled

            # Add section header
            section_header = section_key.replace("_", " ").title()
            para = doc.add_paragraph(section_header)
            style = styles.get("header", {})
            run = para.runs[0]
            run.font.name = style.get("font_name", "Arial")
            run.font.size = Pt(style.get("font_size_pt", 12))
            run.bold = style.get("bold", True)
            run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
            para.paragraph_format.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
            para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
            para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

            # Add section content
            style = styles.get("body", {})
            if section_key == "tables":
                # Handle tables by adding them as actual tables in the doc
                for table_data in section_content:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 1)
                    for row_idx, row in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = style.get("font_name", "Arial")
                                    run.font.size = Pt(style.get("font_size_pt", 11))
                                    run.bold = style.get("bold", False)
                                    run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                                paragraph.paragraph_format.alignment = {
                                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                                }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                                paragraph.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                                paragraph.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))
            elif isinstance(section_content, list):
                # Handle lists (e.g., core competencies, professional experience bullets)
                if section_key == "core_competencies" and style.get("is_horizontal_list", False):
                    # Horizontal list with dots
                    para = doc.add_paragraph(" • ".join(section_content))
                    run = para.runs[0]
                    run.font.name = style.get("font_name", "Arial")
                    run.font.size = Pt(style.get("font_size_pt", 11))
                    run.bold = style.get("bold", False)
                    run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                    para.paragraph_format.alignment = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                    para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                    para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))
                else:
                    # Bullet points
                    for item in section_content:
                        para = doc.add_paragraph(item, style="List Bullet")
                        run = para.runs[0]
                        run.font.name = style.get("font_name", "Arial")
                        run.font.size = Pt(style.get("font_size_pt", 11))
                        run.bold = style.get("bold", False)
                        run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                        para.paragraph_format.alignment = {
                            "left": WD_ALIGN_PARAGRAPH.LEFT,
                            "center": WD_ALIGN_PARAGRAPH.CENTER,
                            "right": WD_ALIGN_PARAGRAPH.RIGHT,
                            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                        }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                        para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                        para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))
            else:
                # Handle paragraphs (e.g., professional summary)
                para = doc.add_paragraph(section_content)
                run = para.runs[0]
                run.font.name = style.get("font_name", "Arial")
                run.font.size = Pt(style.get("font_size_pt", 11))
                run.bold = style.get("bold", False)
                run.font.color.rgb = RGBColor(*style.get("color_rgb", [0, 0, 0]))
                para.paragraph_format.alignment = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }.get(style.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                para.paragraph_format.space_before = Pt(style.get("spacing_before_pt", 6))
                para.paragraph_format.space_after = Pt(style.get("spacing_after_pt", 6))

        # Save the new document to a byte stream
        output_stream = BytesIO()
        doc.save(output_stream)
        output_file = output_stream.getvalue()
        output_stream.close()
        template_stream.close()

        logger.info(f"Created reformatted document (size: {len(output_file)} bytes)")
        return output_file

    except Exception as e:
        logger.error(f"Failed to create reformatted docx: {str(e)}")
        raise