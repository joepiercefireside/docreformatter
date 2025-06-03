from docx import Document
import logging

logger = logging.getLogger(__name__)

def process_docx(file):
    """
    Process a .docx file and extract its content as a string.
    
    Args:
        file: The .docx file stream.
    
    Returns:
        str: The extracted content as a string.
    """
    doc = Document(file)
    content = []
    source_sections = []

    # Extract tables
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_content = [cell.text.strip() for cell in row.cells]
            table_content.append(row_content)
        content.append(f"Table: {table_content}")
        logger.info(f"Extracted table: {table_content}")

    # Extract paragraphs and determine section order
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
            source_sections.append(text)

    logger.info(f"Source section order: {source_sections}")
    return "\n\n".join(content)

def process_text_input(text):
    """
    Process raw text input.
    
    Args:
        text (str): The raw text input.
    
    Returns:
        str: The processed text.
    """
    return text